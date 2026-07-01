"""
Word Document Export Script
============================
Converts a Markdown resume or cover letter to a native Word (.docx) file
using python-docx, which writes real OpenXML — readable by all major ATS parsers.

Usage:
    python scripts/export_docx.py <input.md> <output.docx>

Example:
    python scripts/export_docx.py applications/stripe-biz-analyst/resume-draft.md exports/stripe-biz-analyst/resume.docx
"""

import sys
import re
import os
import subprocess

# ── Dependency check ──────────────────────────────────────────────────────────

def ensure_docx():
    try:
        import docx
    except ImportError:
        print("Installing python-docx...")
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx", "-q"])
        print("Installation complete.\n")

ensure_docx()

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Styling helpers ───────────────────────────────────────────────────────────

def add_horizontal_rule(doc):
    """Add a thin bottom border to the last paragraph (used for section dividers)."""
    para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
    p = para._p
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '000000')
    pBdr.append(bottom)
    pPr.append(pBdr)

def set_font_defaults(doc):
    """Apply Calibri 11pt as the document default."""
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    para_fmt = style.paragraph_format
    para_fmt.space_before = Pt(0)
    para_fmt.space_after = Pt(2)
    para_fmt.line_spacing = Pt(13)

# ── Inline markdown parsing ───────────────────────────────────────────────────

def parse_inline(para, text):
    """Apply bold and italic formatting to a paragraph from markdown inline syntax."""
    # Split on **bold** and *italic* markers
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*)', text)
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = para.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = para.add_run(token[1:-1])
            run.italic = True
        else:
            para.add_run(token)

# ── Document builder ──────────────────────────────────────────────────────────

def build_docx(lines, output_path):
    doc = Document()

    # US Letter dimensions with standard margins
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    set_font_defaults(doc)

    for line in lines:
        line = line.rstrip('\n')

        # H1 — Name (centered, large)
        if line.startswith('# '):
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = para.add_run(line[2:])
            run.bold = True
            run.font.size = Pt(18)
            run.font.name = 'Calibri'

        # H2 — Section heading (uppercase, bordered)
        elif line.startswith('## '):
            para = doc.add_paragraph()
            run = para.add_run(line[3:].upper())
            run.bold = True
            run.font.size = Pt(10)
            run.font.name = 'Calibri'
            add_horizontal_rule(doc)

        # H3 — Role / company line
        elif line.startswith('### '):
            para = doc.add_paragraph()
            run = para.add_run(line[4:])
            run.bold = True
            run.font.size = Pt(11)
            run.font.name = 'Calibri'
            para.paragraph_format.space_before = Pt(4)

        # Horizontal rule (--- in markdown)
        elif re.match(r'^-{3,}$', line.strip()):
            doc.add_paragraph()   # small vertical gap instead of a visible rule

        # Bullet list item
        elif line.startswith('- ') or line.startswith('* '):
            para = doc.add_paragraph(style='List Bullet')
            para.paragraph_format.left_indent = Inches(0.25)
            para.paragraph_format.space_after = Pt(1)
            parse_inline(para, line[2:])

        # Empty line — small spacer
        elif line.strip() == '':
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)

        # Regular paragraph
        else:
            para = doc.add_paragraph()
            parse_inline(para, line)

    os.makedirs(os.path.dirname(output_path) if os.path.dirname(output_path) else '.', exist_ok=True)
    doc.save(output_path)
    print(f"Word document exported to: {output_path}")

# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    if len(sys.argv) != 3:
        print(__doc__)
        sys.exit(1)

    md_path, docx_path = sys.argv[1], sys.argv[2]

    if not os.path.exists(md_path):
        print(f"Error: Input file not found: {md_path}")
        sys.exit(1)

    with open(md_path, encoding='utf-8') as f:
        lines = f.readlines()

    build_docx(lines, docx_path)

if __name__ == "__main__":
    main()
